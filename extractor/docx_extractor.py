from __future__ import annotations

from pathlib import Path

from docx import Document

from extractor.base import BaseExtractor, ExtractionResult


class DocxExtractor(BaseExtractor):
    def extract(self, path: Path) -> ExtractionResult:
        result = ExtractionResult(file_path=str(path))

        try:
            # Open file in binary mode to handle non-ASCII paths
            with open(path, 'rb') as f:
                doc = Document(f)
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

                table_lines = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            table_lines.append(row_text)

                result.text = "\n".join(paragraphs + table_lines)
                result.metadata = {
                    "paragraphs": len(paragraphs),
                    "table_rows": len(table_lines),
                }
            return result
        except Exception as exc:
            result.errors.append(f"DOCX read failed: {exc}")
            return result